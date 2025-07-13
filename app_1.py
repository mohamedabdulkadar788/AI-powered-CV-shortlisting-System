# # cv_shortlisting_app.py

# import streamlit as st
# from sentence_transformers import SentenceTransformer, util
# import pdfplumber
# import os

# # Load model once
# @st.cache_resource
# def load_model():
#     return SentenceTransformer('all-MiniLM-L6-v2')

# model = load_model()

# st.set_page_config(page_title="CV Shortlisting System", layout="wide")
# st.title("📄 Real-time CV Shortlisting System")

# st.markdown("Upload 1 **Job Description** file and up to **5 CVs**. We'll shortlist candidates based on semantic similarity using vector embeddings.")

# # --- Helper Functions ---
# def extract_text_from_pdf(file):
#     with pdfplumber.open(file) as pdf:
#         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# def extract_text(file):
#     if file.name.endswith('.pdf'):
#         return extract_text_from_pdf(file)
#     elif file.name.endswith('.txt'):
#         return file.read().decode('utf-8')
#     else:
#         return ""

# # --- File Uploads ---
# jd_file = st.file_uploader("📌 Upload Job Description (PDF or TXT)", type=['pdf', 'txt'])

# cv_files = st.file_uploader("📥 Upload up to 5 Candidate CVs (PDF or TXT)", type=['pdf', 'txt'], accept_multiple_files=True)

# threshold = st.slider("Set Shortlisting Threshold", min_value=0.5, max_value=0.95, value=0.75)

# if jd_file and cv_files:
#     if len(cv_files) > 5:
#         st.error("Please upload a maximum of 5 CVs.")
#     else:
#         with st.spinner("🔍 Processing..."):
#             jd_text = extract_text(jd_file)
#             jd_embedding = model.encode(jd_text, convert_to_tensor=True)

#             results = []

#             for cv in cv_files:
#                 cv_text = extract_text(cv)
#                 cv_embedding = model.encode(cv_text, convert_to_tensor=True)
#                 score = util.cos_sim(jd_embedding, cv_embedding).item()

#                 status = "✅ Shortlisted" if score >= threshold else "❌ Not Shortlisted"
#                 results.append({
#                     "Candidate": cv.name,
#                     "Similarity Score": round(score, 4),
#                     "Result": status
#                 })

#         # --- Show Results ---
#         st.subheader("📊 Shortlisting Results")
#         for res in results:
#             st.markdown(f"""
#                 **Candidate:** {res['Candidate']}  
#                 **Score:** {res['Similarity Score']}  
#                 **Status:** {res['Result']}  
#                 ---
#             """)


import streamlit as st
from sentence_transformers import SentenceTransformer, util
import pdfplumber
from docx import Document  # <-- using python-docx now

# st.image("proztec.jpg", width=150)

st.set_page_config(page_title="CV Shortlisting System", layout="wide")

col1, col2 = st.columns([1, 6])
with col1:
    st.image("proztec.jpg", width=100)
with col2:
    st.title("📄 CV Shortlisting System - Proz Technologies")


# Load model once
@st.cache_resource
def load_model():
    return SentenceTransformer('all-MiniLM-L6-v2')

model = load_model()

st.set_page_config(page_title="CV Shortlisting System", layout="wide")
st.title("📄 CV Shortlisting System - Proz Technologies")

st.markdown("Upload 1 **Job Description** file and up to **5 CVs**. We'll shortlist candidates based on semantic similarity using vector embeddings.")

# --- Helper Functions ---
def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text(file):
    if file.name.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif file.name.endswith('.txt'):
        return file.read().decode('utf-8')
    elif file.name.endswith('.docx'):
        return extract_text_from_docx(file)
    else:
        return ""

# --- File Uploads ---
jd_file = st.file_uploader("📌 Upload Job Description (PDF, DOCX, or TXT)", type=['pdf', 'docx', 'txt'])

cv_files = st.file_uploader("📥 Upload up to 5 Candidate CVs (PDF, DOCX, or TXT)", type=['pdf', 'docx', 'txt'], accept_multiple_files=True)

threshold = st.slider("Set Shortlisting Threshold", min_value=0.5, max_value=0.95, value=0.75)

if jd_file and cv_files:
    if len(cv_files) > 5:
        st.error("Please upload a maximum of 5 CVs.")
    else:
        with st.spinner("🔍 Processing..."):
            jd_text = extract_text(jd_file)
            jd_embedding = model.encode(jd_text, convert_to_tensor=True)

            results = []

            for cv in cv_files:
                cv_text = extract_text(cv)
                cv_embedding = model.encode(cv_text, convert_to_tensor=True)
                score = util.cos_sim(jd_embedding, cv_embedding).item()

                status = "✅ Shortlisted" if score >= threshold else "❌ Not Shortlisted"
                results.append({
                    "Candidate": cv.name,
                    "Similarity Score": round(score, 4),
                    "Result": status
                })

        # --- Show Results ---
        st.subheader("📊 Shortlisting Results")
        for res in results:
            st.markdown(f"""
                **Candidate:** {res['Candidate']}  
                **Score:** {res['Similarity Score']}  
                **Status:** {res['Result']}  
                ---
            """)




# import streamlit as st
# from sentence_transformers import SentenceTransformer, util
# import pdfplumber
# from docx import Document
# import re
# import nltk
# from nltk.corpus import stopwords

# # --- Download stopwords only once ---
# nltk.download('stopwords')
# stop_words = set(stopwords.words('english'))

# # --- Load model once ---
# @st.cache_resource
# def load_model():
#     return SentenceTransformer('all-MiniLM-L6-v2')

# model = load_model()

# st.set_page_config(page_title="CV Shortlisting System", layout="wide")
# st.title("📄 Real-time CV Shortlisting System")

# st.markdown("Upload 1 **Job Description** file and up to **5 CVs**. We'll shortlist candidates based on semantic similarity using vector embeddings.")

# # --- Text cleaning with stopword removal ---
# def clean_text(text):
#     text = text.lower()  # lowercase
#     text = re.sub(r'\s+', ' ', text)  # remove extra whitespace
#     text = re.sub(r'[^\w\s.,]', '', text)  # remove special characters except basic punctuation
#     words = text.split()
#     filtered = [word for word in words if word not in stop_words]
#     return ' '.join(filtered)

# # --- File extraction functions ---
# def extract_text_from_pdf(file):
#     with pdfplumber.open(file) as pdf:
#         return "\n".join(page.extract_text() or "" for page in pdf.pages)

# def extract_text_from_docx(file):
#     doc = Document(file)
#     return "\n".join([para.text for para in doc.paragraphs])

# def extract_text(file):
#     if file.name.endswith('.pdf'):
#         return extract_text_from_pdf(file)
#     elif file.name.endswith('.txt'):
#         return file.read().decode('utf-8')
#     elif file.name.endswith('.docx'):
#         return extract_text_from_docx(file)
#     else:
#         return ""

# # --- File Uploads ---
# jd_file = st.file_uploader("📌 Upload Job Description (PDF, DOCX, or TXT)", type=['pdf', 'docx', 'txt'])
# cv_files = st.file_uploader("📥 Upload up to 5 Candidate CVs (PDF, DOCX, or TXT)", type=['pdf', 'docx', 'txt'], accept_multiple_files=True)
# threshold = st.slider("Set Shortlisting Threshold", min_value=0.5, max_value=0.95, value=0.75)

# if jd_file and cv_files:
#     if len(cv_files) > 5:
#         st.error("Please upload a maximum of 5 CVs.")
#     else:
#         with st.spinner("🔍 Processing..."):
#             jd_text_raw = extract_text(jd_file)
#             jd_text = clean_text(jd_text_raw)
#             jd_embedding = model.encode(jd_text, convert_to_tensor=True)

#             results = []

#             for cv in cv_files:
#                 cv_text_raw = extract_text(cv)
#                 cv_text = clean_text(cv_text_raw)
#                 cv_embedding = model.encode(cv_text, convert_to_tensor=True)
#                 score = util.cos_sim(jd_embedding, cv_embedding).item()

#                 status = "✅ Shortlisted" if score >= threshold else "❌ Not Shortlisted"
#                 results.append({
#                     "Candidate": cv.name,
#                     "Similarity Score": round(score, 4),
#                     "Result": status
#                 })

#         # --- Show Results ---
#         st.subheader("📊 Shortlisting Results")
#         for res in results:
#             st.markdown(f"""
#                 **Candidate:** {res['Candidate']}  
#                 **Score:** {res['Similarity Score']}  
#                 **Status:** {res['Result']}  
#                 ---
#             """)
